from docx import Document
import os
import re


def _strip_inline_markdown(text: str) -> str:
    """Убирает простейшую инлайн-разметку Markdown (**жирный**, _курсив_, `code`)."""
    if not text:
        return ""
    # Жирный **text** или __text__
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    # Курсив *text* или _text_
    text = re.sub(r"(?<!\*)\*(.+?)\*(?!\*)", r"\1", text)
    text = re.sub(r"(?<!_)_(.+?)_(?!_)", r"\1", text)
    # Инлайн-код `code`
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def _render_markdown_to_docx(doc: Document, content: str) -> None:
    """Очень упрощённый конвертер Markdown → DOCX.

    Поддерживает:
    - заголовки: #, ##, ... → heading уровня 1–4;
    - маркированные списки: строки, начинающиеся с '-', '*', '+' → параграф с буллетом ("• ...");
    - нумерованные списки: '1. text' → обычный параграф с префиксом "1. ";
    - пустые строки → пустые параграфы;
    - инлайн-разметку (**жирный**, _курсив_, `code`) просто очищает от маркеров.
    """
    if not content:
        return

    lines = content.splitlines()

    for raw_line in lines:
        line = raw_line.rstrip("\r\n")
        stripped = line.strip()

        # Пустая строка → визуальный отступ
        if not stripped:
            doc.add_paragraph("")
            continue

        # Заголовки вида '# Текст', '## Текст', ...
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            text = _strip_inline_markdown(heading_match.group(2).strip())
            # В Word уровни 0–4, где 0 — самый крупный заголовок
            doc.add_heading(text, level=level)
            continue

        # Маркированный список '- текст' / '* текст' / '+ текст'
        bullet_match = re.match(r"^[-*+]\s+(.*)$", stripped)
        if bullet_match:
            text = _strip_inline_markdown(bullet_match.group(1).strip())
            doc.add_paragraph(f"• {text}")
            continue

        # Нумерованный список '1. текст' → оставляем как '1. текст'
        numbered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered_match:
            num = numbered_match.group(1)
            text = _strip_inline_markdown(numbered_match.group(2).strip())
            doc.add_paragraph(f"{num}. {text}")
            continue

        # Обычный параграф
        plain_text = _strip_inline_markdown(line)
        doc.add_paragraph(plain_text)


def create_docx(content: str, title: str) -> str:
    """Создает DOCX файл из текста, предварительно очищая простейший Markdown."""
    print(f"Создаю DOCX файл: {title}.docx")
    doc = Document()
    # Внутри документа оставляем полный заголовок, как запросил пользователь
    doc.add_heading(title, 0)

    # Рендерим тело документа из Markdown в человекочитаемый вид
    _render_markdown_to_docx(doc, content)

    if not os.path.exists("output"):
        os.makedirs("output")

    # Делаем безопасное имя файла для файловой системы (Windows/Linux)
    # Разрешаем только буквы/цифры/пробелы/дефис/нижнее подчёркивание, остальное заменяем на "_"
    safe_title = "".join(
        c if (c.isalnum() or c in (" ", "-", "_")) else "_" for c in title
    )
    filename = f"{safe_title.replace(' ', '_')}.docx"

    file_path = os.path.join("output", filename)
    doc.save(file_path)

    print(f"Файл сохранен по пути: {file_path}")
    return file_path